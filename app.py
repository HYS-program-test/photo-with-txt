"""
現場照片報告生成器
------------------
需求套件（requirements.txt）:
    streamlit
    python-docx
    pillow

執行方式：
    streamlit run app.py

拖曳 + 文字敘述說明：
    編號、拖曳目的地、文字敘述欄位這三樣東西，全部放在「同一個」
    st.components.v1.html 的 iframe 裡（同一份 HTML 文件），用 flex 排在同一列，
    這樣藍線一定是滿版寬度，三者也保證對齊，不會再有「兩個元件對不上高度」的問題。

    拖放結果 + 文字敘述，會透過一個隱藏的 st.text_input 同步回 Python：
    JS 把 {"assignments": {...}, "descriptions": {...}} 整理成 JSON，
    用「原生 setter + dispatchEvent + 模擬 Enter/blur」的方式寫進隱藏欄位，
    逼 Streamlit 把值送回後端。這不是官方保證的機制，如果沒反應，
    請看拖曳框最上方那行紅字狀態列，或瀏覽器 DevTools Console 有沒有報錯。

字型：預設嘗試載入「微軟正黑體 msjh.ttc」，若伺服器沒有該字型會退回 Pillow 的可縮放
預設字型（不支援中文）。要讓中文字放大生效，請把中文字型檔案放到跟 app.py 同一層。
"""

import base64
import hashlib
import io
import json
import zipfile
from datetime import datetime

import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

st.set_page_config(page_title="現場照片報告生成器", page_icon="📱", layout="centered")

FONT_PATHS = ["msjh.ttc", "msjhbd.ttc", "TaipeiSansTCBeta-Bold.ttf", "NotoSansTC-Regular.otf"]
ROW_BOX_HEIGHT = 76  # 尚未指派照片的框 / 文字敘述框，統一用這個高度

# -----------------------------------------------------------------------------
# 1. Session State
# -----------------------------------------------------------------------------
if "rows_count" not in st.session_state:
    st.session_state.rows_count = 5
if "assignments" not in st.session_state:
    st.session_state.assignments = {}  # {"01": ["Pxxxxxxxxxx", ...], ...}
if "desc_text" not in st.session_state:
    st.session_state.desc_text = {}  # {"01": "文字敘述...", ...}
if "dnd_sync_last" not in st.session_state:
    st.session_state.dnd_sync_last = ""


def make_photo_key(f):
    h = hashlib.md5(f"{f.name}|{f.size}".encode("utf-8")).hexdigest()[:10]
    return f"P{h}"


# -----------------------------------------------------------------------------
# 2. 照片壓印文字功能：白色 60% 透明底、靠左對齊、微軟正黑體
# -----------------------------------------------------------------------------
def process_photo_with_text(image_bytes, text):
    image = Image.open(io.BytesIO(image_bytes)).convert("RGBA")
    width, height = image.size

    overlay = Image.new("RGBA", image.size, (255, 255, 255, 0))
    draw = ImageDraw.Draw(overlay)

    font_size = max(32, int(min(width, height) * 0.056))
    font = None
    font_used = None
    for path in FONT_PATHS:
        try:
            font = ImageFont.truetype(path, font_size)
            font_used = path
            break
        except IOError:
            continue
    if font is None:
        try:
            font = ImageFont.load_default(size=font_size)
            font_used = "PIL 內建可縮放預設字型（Pillow>=10.1；不支援中文字元）"
        except TypeError:
            font = ImageFont.load_default()
            font_used = "PIL 內建固定大小預設字型（Pillow 版本太舊，不吃 size，也不支援中文）"
    st.session_state["_last_font_used"] = font_used

    margin = int(min(width, height) * 0.025)
    padding = max(10, int(font_size * 0.2))

    lines = text.split("\n") if text else [" "]
    line_heights = [font.getbbox(line)[3] - font.getbbox(line)[1] for line in lines]
    max_line_width = max(font.getbbox(line)[2] - font.getbbox(line)[0] for line in lines)

    box_w = max_line_width + (padding * 2)
    box_h = sum(line_heights) + (padding * 2) + (len(lines) - 1) * 5

    x1 = margin
    y2 = height - margin
    x2 = x1 + box_w
    y1 = y2 - box_h

    draw.rectangle([x1, y1, x2, y2], fill=(255, 255, 255, 153))

    curr_y = y1 + padding
    for line in lines:
        draw.text((x1 + padding, curr_y), line, fill=(0, 0, 0, 255), font=font)
        curr_y += (font.getbbox(line)[3] - font.getbbox(line)[1]) + 5

    combined = Image.alpha_composite(image, overlay)
    output = io.BytesIO()
    combined.convert("RGB").save(output, format="JPEG", quality=95)
    output.seek(0)
    return output


def set_cell_border_none(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        tcBorders.append(border)
    tcPr.append(tcBorders)


# -----------------------------------------------------------------------------
# 3. 產生「拖曳＋文字敘述」合併區的 HTML/JS
#    編號 + 拖曳目的地 + 文字敘述，全部在同一份 HTML、同一列裡，保證對齊
# -----------------------------------------------------------------------------
def build_dnd_html(photo_map, assignments, desc_text, row_keys, box_height):
    assigned_keys = {k for keys in assignments.values() for k in keys}
    pool_keys = [k for k in photo_map if k not in assigned_keys]

    def chip(key):
        b64s = base64.b64encode(photo_map[key].getvalue()).decode()
        return f'<div class="photo-chip" data-photo="{key}"><img src="data:image/jpeg;base64,{b64s}"></div>'

    def esc(s):
        return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    pool_html = "".join(chip(k) for k in pool_keys)

    rows_html = ""
    for rk in row_keys:
        keys = assignments.get(rk, [])
        items_html = "".join(chip(k) for k in keys if k in photo_map)
        desc_val = esc(desc_text.get(rk, ""))
        rows_html += f"""
        <div class="row-block">
            <div class="row-num">{rk}</div>
            <div class="dnd-list row-drop" data-row="{rk}">{items_html}</div>
            <textarea class="row-desc" data-row="{rk}" placeholder="文字敘述...">{desc_val}</textarea>
        </div>"""

    return f"""
    <meta charset="utf-8">
    <style>
      html, body {{ margin:0; padding:0; font-family: "Microsoft JhengHei", -apple-system, sans-serif; }}
      .pool-wrap {{ display:flex; flex-wrap:wrap; gap:8px; padding:8px 4px 14px 4px;
          border-bottom:3px solid #2b7cff; min-height:60px; }}
      .pool-wrap:empty::after {{ content:"（全部照片已指派完畢）"; color:#9aa4b2; font-size:13px; }}
      .photo-chip {{ width:60px; height:60px; border-radius:6px; overflow:hidden;
          border:1px solid #ccc; cursor:grab; flex:none; touch-action:none; }}
      .photo-chip img {{ width:100%; height:100%; object-fit:cover; display:block; pointer-events:none; }}
      .row-block {{ display:flex; align-items:flex-start; gap:10px; border-bottom:1px solid #eee; padding:10px 4px; }}
      .row-num {{ flex:0 0 30px; font-size:20px; font-weight:700; color:#111827; padding-top: calc(({ROW_BOX_HEIGHT}px - 24px)/2); }}
      .row-drop {{ flex:1 1 0; min-width:80px; height:{ROW_BOX_HEIGHT}px; background:#f3f5f8; border-radius:8px;
          padding:6px; display:flex; flex-wrap:wrap; align-content:flex-start; gap:6px;
          touch-action:none; overflow-y:auto; box-sizing:border-box; }}
      .row-drop:empty::after {{ content:"尚未指派照片"; color:#9aa4b2; font-size:11px; }}
      .row-desc {{ flex:1.2 1 0; min-width:100px; height:{ROW_BOX_HEIGHT}px; background:#f3f5f8;
          border:1px solid #e2e5ea; border-radius:8px; padding:8px; font-size:13px;
          font-family:"Microsoft JhengHei", -apple-system, sans-serif; resize:none; box-sizing:border-box; }}
      .sortable-ghost {{ opacity:0.3; }}
      #dnd-status {{ font-size:11px; color:#c0392b; padding:3px 6px; background:#fff8f0;
          border-bottom:1px dashed #eab676; min-height:14px; white-space:pre-wrap; }}
    </style>
    <div id="dnd-status">狀態：載入中…</div>
    <div class="pool-wrap dnd-list" data-row="__pool__">{pool_html}</div>
    <div>{rows_html}</div>
    <script src="https://cdn.jsdelivr.net/npm/sortablejs@1.15.0/Sortable.min.js"></script>
    <script>
    (function () {{
      function log(msg) {{
        var el = document.getElementById('dnd-status');
        if (el) el.textContent = '狀態：' + msg;
      }}
      function findHiddenInput() {{
        try {{
          return window.parent.document.querySelector('input[aria-label="dnd_sync_field"]');
        }} catch (e) {{
          log('錯誤：存取父頁面失敗 - ' + (e && e.message ? e.message : e));
          return null;
        }}
      }}
      function commitToStreamlit(target, jsonStr) {{
        var setter = Object.getOwnPropertyDescriptor(window.parent.HTMLInputElement.prototype, 'value').set;
        setter.call(target, jsonStr);
        target.dispatchEvent(new Event('input', {{ bubbles: true }}));
        target.dispatchEvent(new Event('change', {{ bubbles: true }}));
        target.focus();
        target.dispatchEvent(new KeyboardEvent('keydown', {{ key: 'Enter', code: 'Enter', bubbles: true }}));
        target.dispatchEvent(new KeyboardEvent('keyup', {{ key: 'Enter', code: 'Enter', bubbles: true }}));
        target.blur();
      }}
      function syncState() {{
        var assignments = {{}};
        document.querySelectorAll('.row-drop').forEach(function (el) {{
          var row = el.getAttribute('data-row');
          assignments[row] = Array.from(el.querySelectorAll('.photo-chip')).map(function (c) {{
            return c.getAttribute('data-photo');
          }});
        }});
        var descriptions = {{}};
        document.querySelectorAll('.row-desc').forEach(function (ta) {{
          descriptions[ta.getAttribute('data-row')] = ta.value;
        }});
        var payload = JSON.stringify({{ assignments: assignments, descriptions: descriptions }});
        var target = findHiddenInput();
        if (!target) {{
          log('錯誤：在父頁面找不到隱藏欄位');
          return;
        }}
        try {{
          commitToStreamlit(target, payload);
          log('已同步：' + payload.slice(0, 80));
        }} catch (e) {{
          log('錯誤：' + (e && e.message ? e.message : e));
        }}
      }}
      function initSortable() {{
        if (typeof Sortable === 'undefined') {{
          log('錯誤：Sortable.js 沒有成功載入（可能是網路擋住了 CDN）');
          return;
        }}
        var lists = document.querySelectorAll('.dnd-list');
        lists.forEach(function (el) {{
          new Sortable(el, {{
            group: 'photos',
            animation: 150,
            ghostClass: 'sortable-ghost',
            onEnd: syncState
          }});
        }});
        var timers = {{}};
        document.querySelectorAll('.row-desc').forEach(function (ta) {{
          ta.addEventListener('blur', syncState);
        }});
        log('已初始化，共 ' + lists.length + ' 個拖曳清單，可以開始拖曳／輸入文字');
      }}
      initSortable();
    }})();
    </script>
    """


# -----------------------------------------------------------------------------
# 4. CSS：窄版手機畫面 + 徹底隱藏同步用的 text_input
# -----------------------------------------------------------------------------
st.markdown(
    """
    <style>
    .main .block-container {
        max-width: 550px !important;
        padding-top: 0.5rem !important;
        padding-bottom: 2rem;
    }
    .frozen-header {
        position: sticky;
        top: 0;
        background-color: #ffffff;
        z-index: 9999;
        padding-top: 10px;
        padding-bottom: 6px;
        border-bottom: 3px solid #2b7cff !important;
        margin-bottom: 15px;
    }
    div.st-key-dnd_sync_box {
        position: absolute !important;
        width: 1px !important;
        height: 1px !important;
        overflow: hidden !important;
        opacity: 0.01 !important;
        left: -9999px !important;
        top: -9999px !important;
        margin: 0 !important;
        padding: 0 !important;
    }
    input[aria-label="dnd_sync_field"] {
        opacity: 0.01 !important;
        height: 1px !important;
        width: 1px !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# -----------------------------------------------------------------------------
# 5. 凍結區塊：上傳按鈕（藍色線以上）
# -----------------------------------------------------------------------------
st.markdown('<div class="frozen-header">', unsafe_allow_html=True)

uploaded_files = st.file_uploader(
    "照片上傳 (1~30張)",
    type=["jpg", "jpeg", "png"],
    accept_multiple_files=True,
    key="uploader",
)

photo_map = {}  # {photo_key: UploadedFile, ...}

if uploaded_files:
    if len(uploaded_files) > 30:
        st.warning("最多支援 30 張照片，僅取前 30 張。")
        uploaded_files = uploaded_files[:30]

    for f in uploaded_files:
        photo_map[make_photo_key(f)] = f

    st.caption(f"已上傳 {len(uploaded_files)} 張照片，請在下方拖曳縮圖到各編號列。")

st.markdown("</div>", unsafe_allow_html=True)  # 藍色線結束點

# -----------------------------------------------------------------------------
# 6. 拖曳指派 + 文字敘述（合併在單一 iframe 裡，滿版寬度、同一列）
# -----------------------------------------------------------------------------
if uploaded_files:
    row_keys = [f"{i:02d}" for i in range(1, st.session_state.rows_count + 1)]

    with st.container(key="dnd_sync_box"):
        st.text_input("dnd_sync_field", key="dnd_sync_raw", label_visibility="collapsed")

    raw = st.session_state.get("dnd_sync_raw", "")
    if raw and raw != st.session_state.dnd_sync_last:
        try:
            parsed = json.loads(raw)
            new_assign = parsed.get("assignments", {})
            new_desc = parsed.get("descriptions", {})
            st.session_state.assignments = {
                rk: [p for p in v if p in photo_map]
                for rk, v in new_assign.items()
                if rk in row_keys
            }
            st.session_state.desc_text = {
                rk: v for rk, v in new_desc.items() if rk in row_keys
            }
            st.session_state.dnd_sync_last = raw
        except (json.JSONDecodeError, AttributeError):
            pass

    # 清掉已被移除照片的殘留指派
    pruned = {
        rk: [p for p in v if p in photo_map]
        for rk, v in st.session_state.assignments.items()
    }
    if pruned != st.session_state.assignments:
        st.session_state.assignments = pruned
    assignments = st.session_state.assignments
    desc_text = st.session_state.desc_text

    pool_count = len(photo_map) - sum(len(v) for v in assignments.values())
    pool_lines = max(1, -(-max(pool_count, 1) // 6))
    box_height = 20 + pool_lines * 70 + len(row_keys) * (ROW_BOX_HEIGHT + 22) + 20

    html_code = build_dnd_html(photo_map, assignments, desc_text, row_keys, box_height)
    components.html(html_code, height=box_height, scrolling=False)

    with st.expander("🔧 除錯資訊（測試用，確認沒問題後可以請我刪掉）"):
        st.write("隱藏同步欄位目前收到的原始值 (dnd_sync_raw)：")
        st.code(raw or "(空的，代表 JS 從來沒有成功寫入過)")
        st.write("目前解析後的 assignments：")
        st.json(st.session_state.assignments)
        st.write("目前解析後的 desc_text：")
        st.json(st.session_state.desc_text)
        st.write("目前 photo_map 的 key：")
        st.code(list(photo_map.keys()))
else:
    st.info("請先上傳照片")

if st.button("➕ 新增項目"):
    st.session_state.rows_count += 1
    st.rerun()

st.divider()

# -----------------------------------------------------------------------------
# 7. 匯出下載
# -----------------------------------------------------------------------------
st.subheader("📥 下載專區")
if "_last_font_used" in st.session_state:
    st.caption(f"（上次壓字使用的字型：{st.session_state['_last_font_used']}）")
today_str = datetime.now().strftime("%Y%m%d")

dl_mode = st.radio(
    "照片下載模式：",
    ["📱 手機模式 (逐張下載，方便直接檢視)", "💻 電腦模式 (下載 ZIP 壓縮檔)"],
    horizontal=False,
)

col_dl_photo, col_dl_word = st.columns(2)

with col_dl_photo:
    if dl_mode == "💻 電腦模式 (下載 ZIP 壓縮檔)":
        if st.button("下載照片"):
            if not photo_map or not any(st.session_state.assignments.values()):
                st.warning("請先上傳照片並指派到列表！")
            else:
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                    for idx in range(1, st.session_state.rows_count + 1):
                        row_num = f"{idx:02d}"
                        pkeys = st.session_state.assignments.get(row_num, [])
                        desc = st.session_state.desc_text.get(row_num, "")
                        for n, pkey in enumerate(pkeys, start=1):
                            if pkey not in photo_map:
                                continue
                            img_out = process_photo_with_text(photo_map[pkey].getvalue(), desc)
                            suffix = f"-{n}" if len(pkeys) > 1 else ""
                            zf.writestr(f"{today_str}-{row_num}{suffix}.jpg", img_out.getvalue())
                zip_buffer.seek(0)
                st.download_button(
                    "⬇️ 下載照片 ZIP",
                    data=zip_buffer,
                    file_name=f"Photos_{today_str}.zip",
                    mime="application/zip",
                )
    else:
        for idx in range(1, st.session_state.rows_count + 1):
            row_num = f"{idx:02d}"
            pkeys = st.session_state.assignments.get(row_num, [])
            desc = st.session_state.desc_text.get(row_num, "")
            for n, pkey in enumerate(pkeys, start=1):
                if pkey not in photo_map:
                    continue
                img_out = process_photo_with_text(photo_map[pkey].getvalue(), desc)
                suffix = f"-{n}" if len(pkeys) > 1 else ""
                st.download_button(
                    label=f"⬇️ 下載 {today_str}-{row_num}{suffix}.jpg",
                    data=img_out,
                    file_name=f"{today_str}-{row_num}{suffix}.jpg",
                    mime="image/jpeg",
                    key=f"dl_single_{idx}_{n}",
                )

with col_dl_word:
    if st.button("下載word", use_container_width=True):
        if not photo_map or not any(st.session_state.assignments.values()):
            st.warning("請先上傳照片並指派到列表！")
        else:
            doc = Document()
            for section in doc.sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

            items_to_print = []
            for idx in range(1, st.session_state.rows_count + 1):
                row_num = f"{idx:02d}"
                pkeys = st.session_state.assignments.get(row_num, [])
                desc = st.session_state.desc_text.get(row_num, "")
                for pkey in pkeys:
                    if pkey in photo_map:
                        img_out = process_photo_with_text(photo_map[pkey].getvalue(), desc)
                        items_to_print.append(img_out)

            if not items_to_print:
                st.warning("目前沒有已指派照片的項目！")
            else:
                for page_start in range(0, len(items_to_print), 4):
                    if page_start > 0:
                        doc.add_page_break()

                    page_items = items_to_print[page_start: page_start + 4]
                    table = doc.add_table(rows=2, cols=2)
                    table.autofit = False

                    for cell_idx, img_bytes in enumerate(page_items):
                        r = cell_idx // 2
                        c = cell_idx % 2
                        cell = table.cell(r, c)
                        set_cell_border_none(cell)

                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(img_bytes, width=Inches(3.4))

                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)

                st.download_button(
                    label=f"⬇️ 下載 {today_str}_01.docx",
                    data=doc_io,
                    file_name=f"{today_str}_01.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
